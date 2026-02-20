from typing import List
from io import BytesIO
from docx import Document
from sqlalchemy.orm import Session
import re
from datetime import time

from app.models.discipline import Discipline
from app.models.group import Group
from app.models.user import User, UserRole
from app.models.constraint import TeacherConstraint, ConstraintType


def get_or_create_discipline(db: Session, name: str) -> Discipline:
    d = db.query(Discipline).filter(Discipline.name == name).first()
    if d:
        return d
    d = Discipline(name=name)
    db.add(d)
    db.commit()
    db.refresh(d)
    return d


def get_or_create_group(db: Session, code: str) -> Group:
    g = db.query(Group).filter(Group.code == code).first()
    if g:
        return g
    g = Group(code=code)
    db.add(g)
    db.commit()
    db.refresh(g)
    return g


def get_or_create_teacher(db: Session, full_name: str) -> User:
    u = db.query(User).filter(User.full_name == full_name).first()
    if u:
        return u

    email = (
        full_name.replace(" ", ".")
        .replace(".", "")
        .lower()
        + "@example.com"
    )

    u = User(
        full_name=full_name,
        email=email,
        password_hash="",
        role=UserRole.teacher,
    )
    db.add(u)
    db.commit()
    db.refresh(u)
    return u


def parse_groups_cell(text: str) -> List[str]:
    parts = text.replace(",", " ").split()
    return [p.strip() for p in parts if p.strip()]


def import_assignments_from_docx(db: Session, content: bytes):
    # 🔥 ВАЖНО: открываем DOCX через BytesIO
    file_obj = BytesIO(content)
    doc = Document(file_obj)

    # -----------------------------
    # 1) Таблицы с дисциплинами
    # -----------------------------
    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
        if "наименование дисциплины" not in header_cells[0]:
            continue

        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if not cells or not cells[0]:
                continue

            discipline_name = cells[0]
            groups_cell = cells[1] if len(cells) > 1 else ""

            discipline = get_or_create_discipline(db, discipline_name)
            group_codes = parse_groups_cell(groups_cell)

            for gc in group_codes:
                get_or_create_group(db, gc)

    # -----------------------------
    # 2) Примечания
    # -----------------------------
    notes_started = False
    notes_lines: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if text.startswith("Примечания"):
            notes_started = True
            continue

        if notes_started:
            notes_lines.append(text)

    # -----------------------------
    # 3) Ограничения преподавателей
    # -----------------------------
    fio_pattern = r"[А-ЯЁ][а-яё]+ [А-Я]\.[А-Я]\."

    def add_constraint_all(teachers, ct: ConstraintType):
        for t in teachers:
            c = TeacherConstraint(
                teacher_id=t.id,
                type=ct,
                weekday=None,
                time=None,
                value_bool=None,
            )
            db.add(c)
            db.flush()  # 🔥 предотвращает bulk insert

    for line in notes_lines:
        lower = line.lower()
        names = re.findall(fio_pattern, line)
        teachers: List[User] = [get_or_create_teacher(db, n) for n in names]

        if not teachers:
            continue

        # --- Типовые ограничения ---
        if "не планировать занятия на субботу" in lower or "не ставить занятия на субботу" in lower:
            add_constraint_all(teachers, ConstraintType.no_saturday)

        if "не ставить занятия в пятницу" in lower:
            add_constraint_all(teachers, ConstraintType.no_friday)

        if "после 17.25" in lower or "после 17:25" in lower:
            add_constraint_all(teachers, ConstraintType.no_after_17_25)

        if "не ставить занятия в 8.00" in lower or "не ставить занятия в 8:00" in lower:
            add_constraint_all(teachers, ConstraintType.no_morning)

        if "планировать занятия на субботу" in lower and "не" not in lower:
            add_constraint_all(teachers, ConstraintType.prefer_saturday)

        if "планировать занятия на понедельник" in lower:
            add_constraint_all(teachers, ConstraintType.prefer_monday)

        # --- Фиксированные слоты ---
        if "поставить занятие в среду в 11.40" in lower or "поставить занятие в среду в 11:40" in lower:
            for t in teachers:
                c = TeacherConstraint(
                    teacher_id=t.id,
                    type=ConstraintType.fixed_slot,
                    weekday=3,  # среда
                    time=time(hour=11, minute=40),
                    value_bool=None,
                )
                db.add(c)
                db.flush()

    db.commit()
